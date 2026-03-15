"""
Audio Transcription & Analysis App
- Transcribes audio using AssemblyAI
- Analyzes transcript using Claude (Anthropic API)
- Exports transcript as DOCX
- Exports analyzed attributes as CSV
- Writes analysis results to Google Sheets automatically
- Accepts Exotel webhooks for real-time call processing
"""

import os
import json
import csv
import time
import tempfile
import secrets
import logging
import threading
import requests as http_requests
from datetime import datetime
from pathlib import Path

from flask import Flask, render_template, request, jsonify, send_file
import assemblyai as aai
import anthropic
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Google Sheets imports
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 500 * 1024 * 1024  # 500MB max upload
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", secrets.token_hex(32))

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Use /tmp on Render (ephemeral filesystem)
BASE_TMP = Path(tempfile.gettempdir()) / "audio-analyzer"
UPLOAD_DIR = BASE_TMP / "uploads"
OUTPUT_DIR = BASE_TMP / "outputs"
SESSION_DIR = BASE_TMP / "sessions"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
SESSION_DIR.mkdir(parents=True, exist_ok=True)


# ─── Attribute definitions (must match frontend) ────────────────
INBOUND_ATTRS = [
    "repeat_call_flag", "primary_intent",
    "call_preventable", "secondary_hidden_intent", "intent_recognized_correctly", "complaint_type",
    "complaint_severity", "fcr", "resolution_promised", "unresolved_intent_at_end",
    "overall_sentiment", "sentiment_arc", "emotional_peak_moment", "specific_emotions_detected",
    "emotional_intensity_score", "expressed_satisfaction",
    "agent_mentioned_generic", "customer_reaction_generic", "reason_resistance_generic",
    "generic_switch_completed", "upsell_crosssell_opportunity", "upsell_crosssell_attempt",
    "greeting_quality", "active_listening_signals", "empathy_markers", "closing_quality",
    "agent_handled_objection", "negative_agent_behavior_flags",
]

OUTBOUND_ATTRS = [
    "primary_call_objective", "call_objective_achieved", "reason_for_non_achievement",
    "customer_action_committed", "follow_up_required", "follow_up_to_inbound_complaint",
    "customer_receptivity_at_opening", "opt_out_expressed", "dissatisfaction_with_contact",
    "overall_sentiment", "sentiment_arc", "emotional_peak_moment", "specific_emotions_detected",
    "emotional_intensity_score", "expressed_satisfaction",
    "agent_mentioned_generic", "substitution_pitch_quality", "customer_reaction_generic",
    "reason_resistance_generic", "generic_switch_completed", "upsell_crosssell_opportunity",
    "upsell_crosssell_attempt", "customer_referenced_digital", "competitor_mentioned",
    "greeting_quality", "active_listening_signals", "empathy_markers", "closing_quality",
    "agent_handled_objection", "negative_agent_behavior_flags",
]

META_COLS = ["filename", "call_type", "date", "duration_seconds", "word_count", "transcript"]


def get_attrs_for_type(call_type):
    return OUTBOUND_ATTRS if call_type == "outbound" else INBOUND_ATTRS


# ─── Analysis prompts (used by webhook pipeline) ─────────────────

INBOUND_ANALYSIS_PROMPT = """ROLE
You are an enterprise-grade AI analyst for inbound CRM calls at a healthcare/pharmacy company.
You will receive a pre-generated transcript (produced by an external speech-to-text engine). Your sole task is to analyze it and return structured JSON.

INPUT
- A text transcript of one inbound CRM call (order, query, or complaint).
- Speaker labels (e.g. "Speaker A", "Speaker B") are assigned by the transcription engine. Infer who is the Agent and who is the Customer from conversational context.
- The transcript may contain Hindi, Hinglish (Hindi-English mix), or regional language fragments. Analyze meaning in the original language.
- Transcription artifacts (repeated words, filler sounds, incomplete sentences) are normal.

EVIDENCE STANDARD
- Every analytical value must be directly supported by explicit evidence in the transcript text.
- If evidence is insufficient, ambiguous, or absent, return "Unknown" for that field.
- Never infer, interpolate, speculate, or guess.
- Prefer "Unknown" over incorrect certainty.

OUTPUT FORMAT
Return ONLY a valid JSON object (no markdown fences, no commentary) with exactly two keys:
1. "summary": A 2-3 sentence factual overview of the call.
2. "attributes": An object with EXACTLY these 34 keys in this order:

wait_hold_time_instances, repeat_call_flag, primary_intent, multi_intent_flag, call_preventable, secondary_hidden_intent, intent_recognized_correctly, complaint_type, complaint_severity, fcr, resolution_promised, unresolved_intent_at_end, overall_sentiment, sentiment_arc, emotional_peak_moment, specific_emotions_detected, emotional_intensity_score, closing_sentiment_delta, expressed_satisfaction, agent_mentioned_generic, customer_reaction_generic, reason_resistance_generic, generic_switch_completed, upsell_crosssell_opportunity, upsell_crosssell_attempt, greeting_quality, active_listening_signals, empathy_markers, closing_quality, dead_air_filler_words, agent_handled_objection, negative_agent_behavior_flags, agent_speech_rate, background_noise_agent

STRICT RULES
- Output ONLY the JSON object. No text before or after.
- All values must be flat strings or integers.
- For numeric fields, return integers only.
- Do not add keys beyond the 34 listed above.
- When uncertain, use "Unknown"."""

OUTBOUND_ANALYSIS_PROMPT = """ROLE
You are an enterprise-grade AI analyst for outbound CRM calls at a healthcare/pharmacy company.
You will receive a pre-generated transcript (produced by an external speech-to-text engine). Your sole task is to analyze it and return structured JSON.

INPUT
- A text transcript of one outbound CRM call (follow-up, feedback collection, order confirmation, or proactive outreach).
- Speaker labels are assigned by the transcription engine. Infer who is the Agent and who is the Customer from conversational context.
- The transcript may contain Hindi, Hinglish, or regional language fragments. Analyze meaning in the original language.
- Transcription artifacts are normal.

EVIDENCE STANDARD
- Every analytical value must be directly supported by explicit evidence in the transcript text.
- If evidence is insufficient, ambiguous, or absent, return "Unknown" for that field.
- Never infer, interpolate, speculate, or guess.
- Prefer "Unknown" over incorrect certainty.

OUTPUT FORMAT
Return ONLY a valid JSON object (no markdown fences, no commentary) with exactly two keys:
1. "summary": A 2-3 sentence factual overview of the call.
2. "attributes": An object with EXACTLY these 33 keys in this order:

primary_call_objective, call_objective_achieved, reason_for_non_achievement, customer_action_committed, follow_up_required, follow_up_to_inbound_complaint, customer_receptivity_at_opening, opt_out_expressed, dissatisfaction_with_contact, overall_sentiment, sentiment_arc, emotional_peak_moment, specific_emotions_detected, emotional_intensity_score, closing_sentiment_delta, expressed_satisfaction, agent_mentioned_generic, substitution_pitch_quality, customer_reaction_generic, reason_resistance_generic, generic_switch_completed, upsell_crosssell_opportunity, upsell_crosssell_attempt, customer_referenced_digital, competitor_mentioned, greeting_quality, active_listening_signals, empathy_markers, closing_quality, dead_air_filler_words, agent_handled_objection, negative_agent_behavior_flags, agent_speech_rate

STRICT RULES
- Output ONLY the JSON object. No text before or after.
- All values must be flat strings or integers.
- For numeric fields, return integers only.
- Do not add keys beyond the 33 listed above.
- When uncertain, use "Unknown"."""


def get_analysis_prompt(call_type):
    return OUTBOUND_ANALYSIS_PROMPT if call_type == "outbound" else INBOUND_ANALYSIS_PROMPT


# ─── File-based session store (works across multiple workers) ───

def save_session(session_id, data):
    """Save session data to a JSON file."""
    path = SESSION_DIR / f"{session_id}.json"
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False)


def load_session(session_id):
    """Load session data from a JSON file. Returns None if not found."""
    path = SESSION_DIR / f"{session_id}.json"
    if not path.exists():
        return None
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_clients():
    """Initialize API clients from environment variables."""
    assemblyai_key = os.environ.get("ASSEMBLYAI_API_KEY", "")
    anthropic_key = os.environ.get("ANTHROPIC_API_KEY", "")

    if not assemblyai_key:
        raise ValueError("ASSEMBLYAI_API_KEY environment variable is not set")
    if not anthropic_key:
        raise ValueError("ANTHROPIC_API_KEY environment variable is not set")

    aai.settings.api_key = assemblyai_key
    claude_client = anthropic.Anthropic(api_key=anthropic_key)
    return claude_client


# ═══════════════════════════════════════════════════════════════
# GOOGLE SHEETS INTEGRATION
# ═══════════════════════════════════════════════════════════════

def get_sheets_service():
    """
    Create a Google Sheets API service using a Service Account.

    Requires env var GOOGLE_SERVICE_ACCOUNT_JSON containing the full
    JSON key file contents (as a single string).
    """
    creds_json = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON", "")
    if not creds_json:
        logger.warning("GOOGLE_SERVICE_ACCOUNT_JSON not set — Sheets integration disabled")
        return None

    try:
        creds_data = json.loads(creds_json)
        creds = Credentials.from_service_account_info(
            creds_data,
            scopes=["https://www.googleapis.com/auth/spreadsheets"],
        )
        service = build("sheets", "v4", credentials=creds, cache_discovery=False)
        return service
    except Exception as e:
        logger.error(f"Failed to init Google Sheets service: {e}")
        return None


def get_or_create_sheet_for_type(service, spreadsheet_id, call_type):
    """
    Ensure a sheet tab named 'Inbound' or 'Outbound' exists.
    If it doesn't exist, create it and write headers.
    Returns the sheet name.
    """
    sheet_name = "Inbound" if call_type == "inbound" else "Outbound"
    attrs = get_attrs_for_type(call_type)
    headers = META_COLS + attrs

    try:
        # Check existing sheets
        spreadsheet = service.spreadsheets().get(spreadsheetId=spreadsheet_id).execute()
        existing_sheets = [s["properties"]["title"] for s in spreadsheet.get("sheets", [])]

        if sheet_name not in existing_sheets:
            # Create the sheet tab
            service.spreadsheets().batchUpdate(
                spreadsheetId=spreadsheet_id,
                body={
                    "requests": [{
                        "addSheet": {
                            "properties": {"title": sheet_name}
                        }
                    }]
                },
            ).execute()
            logger.info(f"Created sheet tab: {sheet_name}")

            # Write headers
            service.spreadsheets().values().update(
                spreadsheetId=spreadsheet_id,
                range=f"'{sheet_name}'!A1",
                valueInputOption="RAW",
                body={"values": [headers]},
            ).execute()
            logger.info(f"Wrote {len(headers)} column headers to {sheet_name}")

        return sheet_name

    except Exception as e:
        logger.error(f"Error setting up sheet tab '{sheet_name}': {e}")
        return sheet_name  # return the name anyway; append might still work


def append_to_google_sheet(call_type, metadata, attributes):
    """
    Append one row of analysis results to Google Sheets.

    Env vars needed:
      GOOGLE_SERVICE_ACCOUNT_JSON — full JSON key contents
      GOOGLE_SHEET_ID — the spreadsheet ID from the URL
    """
    spreadsheet_id = os.environ.get("GOOGLE_SHEET_ID", "")
    if not spreadsheet_id:
        logger.info("GOOGLE_SHEET_ID not set — skipping Sheets write")
        return False

    service = get_sheets_service()
    if not service:
        return False

    try:
        sheet_name = get_or_create_sheet_for_type(service, spreadsheet_id, call_type)
        attrs = get_attrs_for_type(call_type)

        # Build ordered row values
        row = []
        for col in META_COLS:
            row.append(str(metadata.get(col, "")))
        for col in attrs:
            row.append(str(attributes.get(col, "Unknown")))

        # Append the row
        service.spreadsheets().values().append(
            spreadsheetId=spreadsheet_id,
            range=f"'{sheet_name}'!A1",
            valueInputOption="RAW",
            insertDataOption="INSERT_ROWS",
            body={"values": [row]},
        ).execute()

        logger.info(f"Appended row to Google Sheet '{sheet_name}' — {metadata.get('filename', 'unknown')}")
        return True

    except Exception as e:
        logger.error(f"Failed to append to Google Sheet: {e}")
        return False


# ═══════════════════════════════════════════════════════════════
# EXOTEL WEBHOOK — Real-time automated pipeline
# ═══════════════════════════════════════════════════════════════

def run_full_pipeline(recording_url, call_sid, call_type, caller_number, duration_seconds):
    """
    Full pipeline: download recording → transcribe → analyze → write to Sheets.
    Runs in a background thread so the webhook returns 200 immediately.
    """
    logger.info(f"[Pipeline] Starting for {call_sid} ({call_type})")

    try:
        # ── Step 1: Initialize API clients ──
        assemblyai_key = os.environ.get("ASSEMBLYAI_API_KEY", "")
        anthropic_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not assemblyai_key or not anthropic_key:
            logger.error("[Pipeline] Missing API keys")
            return

        aai.settings.api_key = assemblyai_key
        claude_client = anthropic.Anthropic(api_key=anthropic_key)

        # ── Step 2: Transcribe with AssemblyAI ──
        # AssemblyAI can transcribe directly from a URL (no download needed)
        logger.info(f"[Pipeline] Transcribing {call_sid}...")
        config = aai.TranscriptionConfig(
            speech_models=["universal-2"],
            speaker_labels=True,
            auto_highlights=True,
        )
        transcriber = aai.Transcriber()
        transcript = transcriber.transcribe(recording_url, config=config)

        if transcript.status == aai.TranscriptStatus.error:
            logger.error(f"[Pipeline] Transcription failed for {call_sid}: {transcript.error}")
            return

        full_text = transcript.text or ""
        if not full_text.strip():
            logger.warning(f"[Pipeline] Empty transcript for {call_sid}, skipping")
            return

        utterances = []
        if transcript.utterances:
            for u in transcript.utterances:
                utterances.append({
                    "speaker": u.speaker,
                    "text": u.text,
                    "start": u.start,
                    "end": u.end,
                })

        actual_duration = transcript.audio_duration or duration_seconds
        word_count = len(full_text.split())
        logger.info(f"[Pipeline] Transcribed {call_sid}: {word_count} words, {actual_duration}s")

        # ── Step 3: Analyze with Claude ──
        logger.info(f"[Pipeline] Analyzing {call_sid} with Claude...")
        analysis_prompt = get_analysis_prompt(call_type)

        system_prompt = """You are an expert audio transcript analyst.
You will be given a transcript and an analysis prompt.
Respond with a JSON object containing two keys:
1. "summary": A free-text analysis based on the prompt.
2. "attributes": An object with key-value pairs of extracted attributes.
   Keys should be short column-friendly names (snake_case).
   Values should be strings or numbers.

Respond ONLY with valid JSON, no markdown fences."""

        user_message = f"""## Transcript
{full_text}

## Analysis Instructions
{analysis_prompt}

Respond with JSON containing "summary" and "attributes" keys."""

        message = claude_client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=4096,
            system=system_prompt,
            messages=[{"role": "user", "content": user_message}],
        )

        response_text = message.content[0].text.strip()

        # Parse JSON (handle markdown fences)
        if response_text.startswith("```"):
            response_text = response_text.split("\n", 1)[1]
            response_text = response_text.rsplit("```", 1)[0]

        result = json.loads(response_text)
        summary = result.get("summary", "")
        attributes = result.get("attributes", {})

        logger.info(f"[Pipeline] Analysis complete for {call_sid}: {len(attributes)} attributes")

        # ── Step 4: Save session (for history view) ──
        session_id = f"exotel_{call_sid}"
        filename = f"{caller_number}_{call_sid}"
        timestamp = datetime.now().isoformat()

        save_session(session_id, {
            "filename": filename,
            "call_type": call_type,
            "full_text": full_text,
            "utterances": utterances,
            "duration_seconds": actual_duration,
            "timestamp": timestamp,
            "caller_number": caller_number,
            "exotel_sid": call_sid,
            "source": "exotel_webhook",
            "analysis": {
                "summary": summary,
                "attributes": attributes,
                "prompt_used": "(auto — webhook pipeline)",
                "model": "claude-sonnet-4-6",
            },
        })

        # ── Step 5: Write to Google Sheets ──
        metadata = {
            "filename": filename,
            "call_type": call_type,
            "date": timestamp[:10],
            "duration_seconds": actual_duration,
            "word_count": word_count,
            "transcript": full_text,
        }
        sheets_ok = append_to_google_sheet(call_type, metadata, attributes)
        if sheets_ok:
            logger.info(f"[Pipeline] Written to Google Sheets for {call_sid}")
        else:
            logger.warning(f"[Pipeline] Sheets write skipped/failed for {call_sid}")

        logger.info(f"[Pipeline] DONE for {call_sid}")

    except json.JSONDecodeError as e:
        logger.error(f"[Pipeline] Claude returned invalid JSON for {call_sid}: {e}")
    except Exception as e:
        logger.error(f"[Pipeline] Error processing {call_sid}: {e}")


@app.route("/api/webhook/exotel", methods=["POST"])
def exotel_webhook():
    """
    Receive call-end notifications from Exotel.

    Exotel sends form-encoded POST data with fields like:
      CallSid, From, To, RecordingUrl, Direction, DialCallDuration, etc.

    Configure this URL in your Exotel "After Call" applet or passthru URL.
    """
    # Verify webhook secret (optional but recommended)
    webhook_secret = os.environ.get("EXOTEL_WEBHOOK_SECRET", "")
    if webhook_secret:
        provided_secret = request.args.get("secret", "") or request.headers.get("X-Webhook-Secret", "")
        if provided_secret != webhook_secret:
            logger.warning("Exotel webhook: invalid secret")
            return jsonify({"error": "unauthorized"}), 401

    # Parse Exotel payload (form-encoded)
    call_sid = request.form.get("CallSid", "") or request.args.get("CallSid", "")
    recording_url = request.form.get("RecordingUrl", "") or request.args.get("RecordingUrl", "")
    direction = (request.form.get("Direction", "") or request.args.get("Direction", "")).lower()
    caller = request.form.get("From", "") or request.args.get("From", "")
    duration_str = request.form.get("DialCallDuration", "0") or request.args.get("DialCallDuration", "0")

    logger.info(f"Exotel webhook received: CallSid={call_sid}, Direction={direction}, RecordingUrl={recording_url}")

    if not recording_url:
        logger.warning(f"Exotel webhook: no RecordingUrl for {call_sid}, skipping")
        return jsonify({"status": "skipped", "reason": "no recording URL"}), 200

    if not call_sid:
        return jsonify({"error": "missing CallSid"}), 400

    # Determine call type from Exotel's Direction field
    # Exotel uses: "inbound" or "outbound" (or "incoming"/"outgoing" in some versions)
    if direction in ("outbound", "outgoing", "out"):
        call_type = "outbound"
    else:
        call_type = "inbound"

    duration_seconds = 0
    try:
        duration_seconds = int(duration_str)
    except (ValueError, TypeError):
        pass

    # Run the full pipeline in a background thread so webhook returns 200 fast
    thread = threading.Thread(
        target=run_full_pipeline,
        args=(recording_url, call_sid, call_type, caller, duration_seconds),
        daemon=True,
    )
    thread.start()

    return jsonify({
        "status": "accepted",
        "call_sid": call_sid,
        "call_type": call_type,
        "message": "Processing started in background",
    }), 200


# ═══════════════════════════════════════════════════════════════
# MANUAL GOOGLE SHEETS WRITE (from the UI)
# ═══════════════════════════════════════════════════════════════

@app.route("/api/sheets/write", methods=["POST"])
def manual_sheets_write():
    """
    Manually push the current analysis to Google Sheets from the UI.
    Expects JSON body: { session_id: "..." }
    """
    data = request.get_json()
    session_id = data.get("session_id")

    session = load_session(session_id) if session_id else None
    if not session:
        return jsonify({"error": "Invalid session."}), 400

    analysis = session.get("analysis")
    if not analysis or not analysis.get("attributes"):
        return jsonify({"error": "No analysis found. Run analysis first."}), 400

    call_type = session.get("call_type", "inbound")
    metadata = {
        "filename": session["filename"],
        "call_type": call_type,
        "date": session["timestamp"][:10],
        "duration_seconds": session.get("duration_seconds", 0),
        "word_count": len(session.get("full_text", "").split()),
        "transcript": session.get("full_text", ""),
    }

    ok = append_to_google_sheet(call_type, metadata, analysis["attributes"])
    if ok:
        return jsonify({"status": "success", "message": "Row added to Google Sheet"})
    else:
        return jsonify({"error": "Failed to write to Google Sheet. Check GOOGLE_SHEET_ID and GOOGLE_SERVICE_ACCOUNT_JSON env vars."}), 500


# ─── Existing Routes ──────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/transcribe", methods=["POST"])
def transcribe():
    """Upload audio and transcribe with AssemblyAI."""
    try:
        get_clients()
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    if "audio" not in request.files:
        return jsonify({"error": "No audio file provided"}), 400

    audio_file = request.files["audio"]
    if audio_file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    # Get call type from form data
    call_type = request.form.get("call_type", "inbound")

    # Save uploaded file
    filename = f"{int(time.time())}_{audio_file.filename}"
    filepath = UPLOAD_DIR / filename
    audio_file.save(filepath)

    try:
        # Configure transcription (universal-2 supports Hindi + multilingual)
        config = aai.TranscriptionConfig(
            speech_models=["universal-2"],
            speaker_labels=True,
            auto_highlights=True,
        )
        transcriber = aai.Transcriber()
        transcript = transcriber.transcribe(str(filepath), config=config)

        if transcript.status == aai.TranscriptStatus.error:
            return jsonify({"error": f"Transcription failed: {transcript.error}"}), 500

        # Build structured transcript
        utterances = []
        if transcript.utterances:
            for u in transcript.utterances:
                utterances.append({
                    "speaker": u.speaker,
                    "text": u.text,
                    "start": u.start,
                    "end": u.end,
                })

        full_text = transcript.text or ""

        # Store in session (file-based for multi-worker support)
        session_id = filename.split(".")[0]
        save_session(session_id, {
            "filename": audio_file.filename,
            "call_type": call_type,
            "full_text": full_text,
            "utterances": utterances,
            "duration_seconds": (transcript.audio_duration or 0),
            "timestamp": datetime.now().isoformat(),
        })

        return jsonify({
            "session_id": session_id,
            "text": full_text,
            "utterances": utterances,
            "duration_seconds": transcript.audio_duration,
            "word_count": len(full_text.split()),
        })

    except Exception as e:
        return jsonify({"error": f"Transcription error: {str(e)}"}), 500
    finally:
        # Cleanup uploaded file
        if filepath.exists():
            filepath.unlink()


@app.route("/api/analyze", methods=["POST"])
def analyze():
    """Analyze transcript with Claude using a custom prompt."""
    try:
        claude_client = get_clients()
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    data = request.get_json()
    session_id = data.get("session_id")
    analysis_prompt = data.get("prompt", "")
    model = data.get("model", "claude-sonnet-4-6")

    session = load_session(session_id) if session_id else None
    if not session:
        return jsonify({"error": "Invalid session. Please transcribe audio first."}), 400
    if not analysis_prompt.strip():
        return jsonify({"error": "Analysis prompt cannot be empty."}), 400
    transcript_text = session["full_text"]

    # Build the message for Claude
    system_prompt = """You are an expert audio transcript analyst.
You will be given a transcript and an analysis prompt.
Respond with a JSON object containing two keys:
1. "summary": A free-text analysis based on the prompt.
2. "attributes": An object with key-value pairs of extracted attributes.
   Keys should be short column-friendly names (snake_case).
   Values should be strings or numbers.

Respond ONLY with valid JSON, no markdown fences."""

    user_message = f"""## Transcript
{transcript_text}

## Analysis Instructions
{analysis_prompt}

Respond with JSON containing "summary" and "attributes" keys."""

    try:
        message = claude_client.messages.create(
            model=model,
            max_tokens=4096,
            system=system_prompt,
            messages=[{"role": "user", "content": user_message}],
        )

        response_text = message.content[0].text.strip()

        # Parse JSON response (handle markdown fences if present)
        if response_text.startswith("```"):
            response_text = response_text.split("\n", 1)[1]
            response_text = response_text.rsplit("```", 1)[0]

        result = json.loads(response_text)
        summary = result.get("summary", "")
        attributes = result.get("attributes", {})

        # Store analysis in session
        session["analysis"] = {
            "summary": summary,
            "attributes": attributes,
            "prompt_used": analysis_prompt,
            "model": model,
        }
        save_session(session_id, session)

        # Auto-write to Google Sheets (if configured)
        call_type = session.get("call_type", "inbound")
        metadata = {
            "filename": session["filename"],
            "call_type": call_type,
            "duration_seconds": session.get("duration_seconds", 0),
            "word_count": len(session["full_text"].split()),
            "date": session["timestamp"][:10],
            "transcript": session["full_text"],
        }
        sheets_ok = append_to_google_sheet(call_type, metadata, attributes)

        return jsonify({
            "summary": summary,
            "attributes": attributes,
            "metadata": metadata,
            "sheets_written": sheets_ok,
        })

    except json.JSONDecodeError:
        # If Claude didn't return valid JSON, return the raw text
        session["analysis"] = {
            "summary": response_text,
            "attributes": {},
            "prompt_used": analysis_prompt,
            "model": model,
        }
        save_session(session_id, session)
        return jsonify({
            "summary": response_text,
            "attributes": {},
            "note": "Claude did not return structured JSON. Try refining your prompt."
        })
    except Exception as e:
        return jsonify({"error": f"Analysis error: {str(e)}"}), 500


@app.route("/api/export/docx", methods=["POST"])
def export_docx():
    """Export transcript as a formatted DOCX file."""
    data = request.get_json()
    session_id = data.get("session_id")

    session = load_session(session_id) if session_id else None
    if not session:
        return jsonify({"error": "Invalid session."}), 400

    doc = Document()

    # --- Title ---
    title = doc.add_heading("Audio Transcript", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Metadata ---
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"File: {session['filename']}  |  Date: {session['timestamp'][:10]}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    duration = session.get("duration_seconds", 0)
    if duration:
        minutes = int(duration // 60)
        seconds = int(duration % 60)
        dur_run = meta.add_run(f"  |  Duration: {minutes}m {seconds}s")
        dur_run.font.size = Pt(10)
        dur_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("")  # spacer

    # --- Transcript body ---
    doc.add_heading("Transcript", level=1)

    utterances = session.get("utterances", [])
    if utterances:
        for u in utterances:
            p = doc.add_paragraph()
            speaker_run = p.add_run(f"Speaker {u['speaker']}: ")
            speaker_run.bold = True
            speaker_run.font.size = Pt(11)
            speaker_run.font.color.rgb = RGBColor(44, 62, 80)
            text_run = p.add_run(u["text"])
            text_run.font.size = Pt(11)
    else:
        p = doc.add_paragraph(session["full_text"])
        p.style.font.size = Pt(11)

    # --- Analysis section (if available) ---
    analysis = session.get("analysis")
    if analysis:
        doc.add_page_break()
        doc.add_heading("Analysis", level=1)

        prompt_p = doc.add_paragraph()
        prompt_label = prompt_p.add_run("Prompt used: ")
        prompt_label.bold = True
        prompt_label.font.size = Pt(10)
        prompt_text = prompt_p.add_run(analysis["prompt_used"])
        prompt_text.font.size = Pt(10)
        prompt_text.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph("")
        doc.add_heading("Summary", level=2)
        summary_p = doc.add_paragraph(analysis["summary"])
        summary_p.style.font.size = Pt(11)

        if analysis.get("attributes"):
            doc.add_heading("Extracted Attributes", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Attribute"
            hdr[1].text = "Value"
            for key, val in analysis["attributes"].items():
                row = table.add_row().cells
                row[0].text = str(key)
                row[1].text = str(val)

    # Save
    output_path = OUTPUT_DIR / f"transcript_{session_id}.docx"
    doc.save(str(output_path))

    return send_file(
        str(output_path),
        as_attachment=True,
        download_name=f"transcript_{session['filename'].rsplit('.', 1)[0]}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/api/export/csv", methods=["POST"])
def export_csv():
    """Export analyzed attributes as a CSV row."""
    data = request.get_json()
    session_id = data.get("session_id")

    session = load_session(session_id) if session_id else None
    if not session:
        return jsonify({"error": "Invalid session."}), 400

    analysis = session.get("analysis")

    if not analysis or not analysis.get("attributes"):
        return jsonify({"error": "No analysis attributes found. Run analysis first."}), 400

    attributes = analysis["attributes"]

    # Add metadata columns
    row = {
        "filename": session["filename"],
        "date": session["timestamp"][:10],
        "duration_seconds": session.get("duration_seconds", ""),
        "word_count": len(session["full_text"].split()),
        **attributes,
    }

    output_path = OUTPUT_DIR / f"analysis_{session_id}.csv"
    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=row.keys())
        writer.writeheader()
        writer.writerow(row)

    return send_file(
        str(output_path),
        as_attachment=True,
        download_name=f"analysis_{session['filename'].rsplit('.', 1)[0]}.csv",
        mimetype="text/csv",
    )


@app.route("/api/sessions", methods=["GET"])
def list_sessions():
    """List all sessions with full data for history view."""
    sessions = []
    for path in SESSION_DIR.glob("*.json"):
        sid = path.stem
        data = load_session(sid)
        if data:
            analysis = data.get("analysis", {})
            duration = data.get("duration_seconds", 0)
            entry = {
                "session_id": sid,
                "filename": data["filename"],
                "call_type": data.get("call_type", "inbound"),
                "timestamp": data["timestamp"],
                "duration_seconds": duration,
                "word_count": len(data.get("full_text", "").split()),
                "full_text": data.get("full_text", ""),
                "has_analysis": bool(analysis),
                "attributes": analysis.get("attributes", {}),
                "summary": analysis.get("summary", ""),
            }
            sessions.append(entry)
    # Sort by timestamp descending (most recent first)
    sessions.sort(key=lambda x: x["timestamp"], reverse=True)
    return jsonify(sessions)


@app.route("/api/health", methods=["GET"])
def health():
    """Health check endpoint for Render."""
    has_aai = bool(os.environ.get("ASSEMBLYAI_API_KEY"))
    has_ant = bool(os.environ.get("ANTHROPIC_API_KEY"))
    has_sheets = bool(os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON")) and bool(os.environ.get("GOOGLE_SHEET_ID"))
    has_exotel = bool(os.environ.get("EXOTEL_WEBHOOK_SECRET"))
    return jsonify({
        "status": "ok",
        "assemblyai_configured": has_aai,
        "anthropic_configured": has_ant,
        "google_sheets_configured": has_sheets,
        "exotel_webhook_configured": has_exotel,
    })


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("FLASK_ENV") != "production"
    print(f"\n  Audio Transcription & Analysis App")
    print(f"  Running on http://0.0.0.0:{port}")
    print(f"  Debug: {debug}\n")
    app.run(debug=debug, host="0.0.0.0", port=port)
